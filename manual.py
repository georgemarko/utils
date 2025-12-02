import requests
import shutil
import zipfile
import os
import csv
from lxml import etree
from docx import Document
from copy import deepcopy
from collections import Counter
from credentials import MYUSERNAME, MYPASSWORD
#from docx2pdf import convert

############## CONFIGURATION FLAGS ################################

# Set to True for SEEMP I-II
# Set to False for SEEMP PART III
SEEMP_VERSION_1_2 = True

# Include waste incinerator in tables
INCLUDE_WASTE_INCINERATOR = True

######################################################

############## INPUTS ################################

# If SELECT_ONE_IMO is True then only the VESSEL_IMO is used

SELECT_ONE_IMO = True
VESSEL_IMO = "9543691"
COMPANY_NAME = "COMPANY"

######################################################


# ---------------- CONFIG ----------------
API_BASE_URL = "https://mariner.alphamrn.com/api"
EMISSION_SOURCES_URL = f"{API_BASE_URL}/emission-sources"
CSV_FILE = "vessels.csv"
WORD_TEMPLATE = "model.docx" if SEEMP_VERSION_1_2 else "model_3.docx"
TEMP_DIR = "temp_docx_extract"
RESULTS_DIR = "results"

if not os.path.exists(RESULTS_DIR):
    os.makedirs(RESULTS_DIR)




def authenticate(username, password, base_url=API_BASE_URL, remember_me=False):
    url = f"{base_url}/authenticate"
    
    payload = {
        "username": username,
        "password": password,
        "rememberMe": remember_me
    }
    
    response = requests.post(url, json=payload)
    
    if response.status_code == 200:
        data = response.json()
        jwt_token = data.get("id_token")
        return jwt_token
    else:
        print(f"Authentication failed: {response.status_code}")
        print(response.text)
        return None

token = authenticate(MYUSERNAME, MYPASSWORD)

HEADERS = {
    "Authorization": f"Bearer {token}",
    "Accept": "application/json"
}


# ---------------- PLACEHOLDER MAPPING ----------------
PLACEHOLDER_MAP = {
    "{{VSLNAME}}": "vesselName",
    "{{IMO}}": "imo",
    "{{DWT}}": "deadWeight",
    "{{DWTVALUE}}": "deadWeightValue",
    "{{HULL}}": "hullNo",
    "{{VSLTYPE}}": "vesselType",
    "{{VSLTYPENAME}}": "vesselTypeName",
    "{{DWG}}": "dwg",  # From CSV
    "{{COUNTRY}}": "flagCountryName",
    "{{PORT}}": "registryPort",
    "{{CALLSIGN}}": "callsign",
    "{{GROSSTONNAGE}}": "grossTonnage",
    "{{NETTONNAGE}}": "netTonnage",
    "{{EEDI}}": "aEedi",
    "{{EEXI}}": "aEexi",
    "{{ICECLASS}}": "iceClass",
    "{{BUILDER}}": "shipbuilder",
    "{{YEAR}}": "deliveryYear",
    "{{LENGTHOA}}": "overallLength",
    "{{LENGTHBP}}": "lengthBp",
    "{{BREADTH}}": "breadth",
    "{{DEPTH}}": "depth",
    "{{SLD}}": "summerLoadDraught"
}

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

# ---------------- DOCX UTILITIES ----------------
def replace_placeholders_in_t_nodes(t_nodes, placeholders):
    """
    Replace placeholders inside each <w:t> node individually.
    Handles multiple placeholders in the same run and surrounding static text.
    """
    for t in t_nodes:
        if t.text:
            for ph, replacement in placeholders.items():
                if ph in t.text:
                    t.text = t.text.replace(ph, str(replacement))
    return True


def recursive_replace(element, placeholders, visited=None):
    """
    Recursively replace all placeholders in Word elements,
    including tables, text boxes, content controls, headers, footers, etc.
    """
    if visited is None:
        visited = set()

    element_id = id(element)
    if element_id in visited:
        return False
    visited.add(element_id)

    replaced = False

    t_nodes = list(element.iter("{%s}t" % W_NS))
    if t_nodes:
        if replace_placeholders_in_t_nodes(t_nodes, placeholders):
            replaced = True

    for drawing in element.iter("{%s}drawing" % W_NS):
        for txbx in drawing.iter("{http://schemas.microsoft.com/office/word/2010/wordprocessingShape}txbxContent"):
            if recursive_replace(txbx, placeholders, visited):
                replaced = True

    for sdt in element.iter("{%s}sdtContent" % W_NS):
        if sdt is not element:
            if recursive_replace(sdt, placeholders, visited):
                replaced = True

    return replaced



def process_docx(input_path, output_path, placeholders):
    if os.path.exists(TEMP_DIR):
        shutil.rmtree(TEMP_DIR)
    os.makedirs(TEMP_DIR)
    with zipfile.ZipFile(input_path, 'r') as zip_ref:
        zip_ref.extractall(TEMP_DIR)
    word_dir = os.path.join(TEMP_DIR, "word")
    for root_dir, dirs, files in os.walk(word_dir):
        for file in files:
            if file.endswith(".xml"):
                file_path = os.path.join(root_dir, file)
                parser = etree.XMLParser(remove_blank_text=False)
                tree = etree.parse(file_path, parser)
                root = tree.getroot()
                recursive_replace(root, placeholders)
                tree.write(file_path, encoding="UTF-8", xml_declaration=True)
    with zipfile.ZipFile(output_path, "w") as zipf:
        for root_dir, dirs, files in os.walk(TEMP_DIR):
            for file in files:
                file_path = os.path.join(root_dir, file)
                arcname = os.path.relpath(file_path, TEMP_DIR)
                zipf.write(file_path, arcname)
    shutil.rmtree(TEMP_DIR)

# ---------------- CSV UTILITIES ----------------
def get_imos_for_company(csv_file, company_name=None, vessel_imo=None):
    imos = {}
    with open(csv_file, newline="", encoding="utf-8") as f:
        reader = csv.DictReader(f)
        for row in reader:
            if vessel_imo != None:
                if row["IMO"] == f"{vessel_imo}":
                    imos[row["IMO"]] = row
            elif company_name != None:
                if row["COMPANY NAME"] == company_name:
                    imos[row["IMO"]] = row
    return imos

# ---------------- API FETCH ----------------
def get_vessel(imo):
    resp = requests.get(f"{API_BASE_URL}/vessels/imo/{imo}", headers=HEADERS)
    resp.raise_for_status()
    return resp.json()

def get_emission_sources_for_imo(imo):
    resp = requests.get(f"{EMISSION_SOURCES_URL}/vessel/{imo}", headers=HEADERS)
    resp.raise_for_status()
    return resp.json()

def format_number(value):
    """Helper to format numeric values with thousand separators"""
    try:
        return f"{float(value):,}"
    except (TypeError, ValueError):
        return "N/A"

# ---------------- PLACEHOLDER HELPERS ----------------
def format_vessel_placeholder(vessel, csv_dwg):
    result = {}
    for ph, field in PLACEHOLDER_MAP.items():
        if field == "dwg":
            value = csv_dwg
        elif field == "vesselTypeName":
            value = vessel.get("vesselType", "")
            value = " ".join(word.upper() for word in value.split('_'))
        elif field == "vesselType":
            value = vessel.get("vesselType", "")
            value = " ".join(word.capitalize() for word in value.split('_'))
        elif field == "aEedi":
            eedi = vessel.get("aEedi")
            value = f"{format_number(eedi)} gr CO₂ / ton-mile" if eedi else "N/A"

        elif field == "aEexi":
            eexi = vessel.get("aEexi")
            value = f"{format_number(eexi)} gr CO₂ / ton-mile" if eexi else "N/A"

        elif field in ["grossTonnage", "netTonnage"]:
            num = vessel.get(field)
            if num is not None:
                value = f"{num:,.2f}".rstrip("0").rstrip(".")
            else:
                value = "N/A"
        elif field == "deadWeightValue":
            num = vessel.get("deadWeight")
            if num is not None:
                value = f"{int(num):,}" if num == int(num) else f"{num:,}"
            else:
                value = "N/A"
        elif field == "deadWeight":
            num = vessel.get(field)
            if num is not None:
                val = f"{int(num):,}" if num == int(num) else f"{num:,}"
                value = f"{val} MT"
            else:
                value = "N/A"
        elif field in ["overallLength", "lengthBp", "breadth", "depth", "summerLoadDraught"]:
            num = vessel.get(field)
            if num is not None:
                val = f"{num:,.2f}".rstrip(".")
                value = f"{val} m"
            else:
                value = "N/A"
        elif "." in field:
            parts = field.split(".")
            val = vessel
            for part in parts:
                val = val.get(part, {})
            value = val if isinstance(val, str) else str(val)
        elif field == "deliveryYear":
            value = vessel.get("deliveryYear", "N/A")
        else:
            value = vessel.get(field, "N/A")

        if value is None:
            value = "N/A"
        result[ph] = str(value)
    return result

def format_fuel_types(emission_sources, include_bio):
    rows = []
    custom_order = ["Main Engine", "Auxiliary Engine", "Fired Boiler", "Inert Gas Generator", "Waste Incinerator"]
    for source in emission_sources:
        bio_value = "Biofuels"
        type_name = source.get("type")
        
        if "waste incinerator" in type_name.lower() and not INCLUDE_WASTE_INCINERATOR:
            continue
        
        if "boiler" in type_name.lower():
            row = {"TYPE": "Fired Boiler",
                   "HFO": "HFO",
                   "LFO": "LFO",
                   "MGO": "MGO / MDO"}
        elif type_name.lower() in ["inert gas generator", "waste incinerator"]:
            row = {"TYPE": type_name,
                   "HFO": "",
                   "LFO": "",
                   "MGO": "MGO / MDO"}
            bio_value = ""
        else:
            row = {
            "TYPE": type_name,
            "HFO": "HFO",
            "LFO": "LFO",
            "MGO": "MGO / MDO"
        }
        if include_bio:
            row["BIO"] = bio_value
        rows.append(row)
    data = sorted(rows, key=lambda x: custom_order.index(x['TYPE']) if x['TYPE'] in custom_order else len(custom_order))
    type_counts = Counter(item['TYPE'] for item in data)
    type_numbering = {es_type: 0 for es_type, count in type_counts.items() if count > 1}

    for item in data:
        es_type = item['TYPE']
        if es_type in type_numbering:
            type_numbering[es_type] += 1
            item['TYPE'] = f"{es_type} No. {type_numbering[es_type]}"

    return data

ENGINE_CONFIG = {
    "main engine": {"cylinders": 6, "stroke": 2},
    "auxiliary engine": {"cylinders": 6, "stroke": 4},
    "hydraulic power pack": {"cylinders": 6, "stroke": 4}
}

import re

def alphanumeric_key(s):
    return [int(c) if c.isdigit() else c.lower() for c in re.split('([0-9]+)', s)]

def format_emission_sources(emission_sources, verifier=None):
    lines = []
    source_type_order = ['main engine', 'auxiliary engine', 'boiler', 'inert gas generator', 'waste incinerator']
    
    normalized_sources = []
    for s in emission_sources:
        original_type = s.get("type", "Unknown").lower()
        
        if "waste incinerator" in original_type and not INCLUDE_WASTE_INCINERATOR:
            continue
        
        if "boiler" in original_type:
            normalized_type = "Fired Boiler"
        elif "hydraulic power pack" in original_type:
            normalized_type = "Auxiliary Engine"
        else:
            normalized_type = s.get("type", "Unknown")
        
        normalized_sources.append({
            "source": s,
            "normalized_type": normalized_type,
            "original_type": original_type
        })
    
    type_counts = Counter(item['normalized_type'] for item in normalized_sources)
    type_counters = {t: 0 for t in type_counts}
    
    def sort_key(item):
        nt = item['normalized_type'].lower()
        for idx, target in enumerate(source_type_order):
            if target in nt:
                return (idx, alphanumeric_key(item['source'].get('identificationNumber', '')))
        return (len(source_type_order), alphanumeric_key(item['source'].get('identificationNumber', '')))
    
    normalized_sources.sort(key=sort_key)
    
    for item in normalized_sources:
        source = item['source']
        normalized_type = item['normalized_type']
        original_type = item['original_type']
        
        name = normalized_type
        
        if type_counts[normalized_type] > 1:
            type_counters[normalized_type] += 1
            name += f" No. {type_counters[normalized_type]}"
        
        if source.get("manufacturer") and source.get("model"):
            name += f" {source['manufacturer']} {source['model']}"
        elif source.get("manufacturer"):
            name += f" {source['manufacturer']}"
        elif source.get("model"):
            name += f" {source['model']}"

        parts = []
        
        # Check if technical description exists for boilers (after stripping whitespace)
        tech_desc = source.get("technicalDescription", "").strip() if "boiler" in original_type and source.get("technicalDescription") else ""
        
        if tech_desc:
            # Use technical description as the whole text for boilers
            details = tech_desc
        elif "boiler" in original_type:
            # For boilers without technical description
            rp = source.get("ratingPowerValue")
            rpu = source.get("ratingPowerUnit", "")
            if rp:
                parts.append(f"Capacity {rp} {rpu}".strip())
            parts.append("Oil fired boiler")
            
            rpm = source.get("rpm")
            if rpm:
                parts.append(f"at {rpm} RPM")
            sfocv = source.get("sfocValue")
            sfocmax = source.get("sfocMaxValue")
            sfocunit = source.get("sfocUnit", "")
            if sfocv:
                foc_label = "FOC"
                sfoc_text = f"{foc_label} {sfocv}"
                if sfocmax:
                    sfoc_text += f"-{sfocmax}"
                if sfocunit:
                    sfoc_text += f" {sfocunit}"
                parts.append(sfoc_text)
            year = source.get("yearOfInstallation")
            if year:
                parts.append(f"Installation Year {year}")
            serial = source.get("identificationNumber")
            if serial:
                parts.append(f"Serial No. {serial}")
            
            details = ", ".join(parts)
        else:
            # For other engine types
            rp = source.get("ratingPowerValue") 
            rpu = source.get("ratingPowerUnit", "")
            if rp:
                parts.append(f"{rp} {rpu}".strip())
            
            rpm = source.get("rpm")
            if rpm:
                parts.append(f"at {rpm} RPM")
            sfocv = source.get("sfocValue")
            sfocmax = source.get("sfocMaxValue")
            sfocunit = source.get("sfocUnit", "")
            if sfocv:
                foc_label = "FOC" if "waste incinerator" in original_type or "inert gas generator" in original_type else "SFOC"
                sfoc_text = f"{foc_label} {sfocv}"
                if sfocmax:
                    sfoc_text += f"-{sfocmax}"
                if sfocunit:
                    sfoc_text += f" {sfocunit}"

                try:
                    verifier_val = verifier.strip().lower() if isinstance(verifier, str) and verifier is not None else ""
                except Exception:
                    verifier_val = ""
                if "auxiliary" in normalized_type.lower():
                    mcr_note = "at 50% MCR" if verifier_val == "rina" else "at 100% MCR"
                    sfoc_text = f"{sfoc_text} {mcr_note}"

                parts.append(sfoc_text)
            year = source.get("yearOfInstallation")
            if year:
                parts.append(f"Installation Year {year}")
            serial = source.get("identificationNumber")
            if serial:
                parts.append(f"Serial No. {serial}")

            if original_type in ENGINE_CONFIG:
                cfg = ENGINE_CONFIG[original_type]
                parts.append(f"{cfg['cylinders']}-cylinder, {cfg['stroke']}-stroke")

            details = ", ".join(parts)

        lines.append({"MODEL": name, "DETAILS": details})
    
    return lines

# ---------------- TABLE UTILITIES ----------------
def add_text_with_superscript(paragraph, text, base_run):
    """
    Add text to paragraph, converting ^ notation to superscript.
    E.g., "m^3/h" becomes "m³/h" with 3 as superscript.
    """
    if '^' not in text:
        base_run.text = text
        return
    
    base_run.text = ""
    
    parts = text.split('^')
    for idx, part in enumerate(parts):
        if idx == 0:
            new_run = paragraph.add_run(part)
            new_run.font.name = base_run.font.name
            new_run.font.size = base_run.font.size
            new_run.bold = base_run.bold
            new_run.italic = base_run.italic
        else:
            if len(part) > 0:
                super_run = paragraph.add_run(part[0])
                super_run.font.superscript = True
                super_run.font.name = base_run.font.name
                super_run.font.size = base_run.font.size
                super_run.bold = base_run.bold
                super_run.italic = base_run.italic
                
                if len(part) > 1:
                    normal_run = paragraph.add_run(part[1:])
                    normal_run.font.name = base_run.font.name
                    normal_run.font.size = base_run.font.size
                    normal_run.bold = base_run.bold
                    normal_run.italic = base_run.italic

def replace_placeholder_preserve_format(paragraph, replacements):
    runs = paragraph.runs
    i = 0
    while i < len(runs):
        run = runs[i]
        text_buffer = run.text
        j = i + 1
        while j < len(runs) and any(f"{{{{{key}}}}}" not in text_buffer for key in replacements):
            text_buffer += runs[j].text
            j += 1
        for key, val in replacements.items():
            if f"{{{{{key}}}}}" in text_buffer:
                text_buffer = text_buffer.replace(f"{{{{{key}}}}}", val)
                add_text_with_superscript(paragraph, text_buffer, run)
                for k in range(i+1, j):
                    runs[k].text = ""
                i = j - 1
                break
        i += 1

def populate_table(doc, data_rows, placeholders):
    target_table = None
    template_row = None
    for table in doc.tables:
        for row in table.rows:
            row_text = " ".join(cell.text for cell in row.cells)
            if all(ph in row_text for ph in placeholders):
                target_table = table
                template_row = row
                break
        if template_row:
            break
    if target_table is None:
        raise ValueError("No table found containing placeholders.")
    for data in data_rows:
        new_tr = deepcopy(template_row._tr)
        target_table._tbl.append(new_tr)
        new_row = target_table.rows[-1]
        for cell in new_row.cells:
            for paragraph in cell.paragraphs:
                replace_placeholder_preserve_format(paragraph, data)
    target_table._tbl.remove(template_row._tr)


def has_bio(doc):
    target_table = None
    template_row = None
    include_bio = False

    for table in doc.tables:
        for row in table.rows:
            row_text = "".join(cell.text for cell in row.cells).replace(" ", "").replace("\n", "")
            required_ph = ["{{TYPE}}", "{{HFO}}", "{{LFO}}", "{{MGO}}"]
            if all(ph.replace(" ", "") in row_text for ph in required_ph):
                target_table = table
                template_row = row
                if "{{BIO}}" in row_text:
                    include_bio = True
                break
        if template_row:
            break

    return include_bio

def get_method_from_placeholder(doc):
    """
    Find the table cell containing {{METHOD}} and return the text
    from the cell above it.
    """
    for table in doc.tables:
        for row_idx, row in enumerate(table.rows):
            for cell_idx, cell in enumerate(row.cells):
                if "{{METHOD}}" in cell.text:
                    if row_idx > 0:
                        return table.rows[row_idx - 1].cells[cell_idx].text.strip()
                    else:
                        return ""
    return ""


def get_issue_number(doc):
    """
    Find the table with 'Issue Number' text and return the last cell value
    from that column.
    """
    for table in doc.tables:
        issue_col_idx = None
        for row in table.rows:
            for cell_idx, cell in enumerate(row.cells):
                cell_text = cell.text.strip()
                if "Issue Number" in cell_text or "Issue No" in cell_text:
                    issue_col_idx = cell_idx
                    break
            if issue_col_idx is not None:
                break
        
        if issue_col_idx is not None:
            for row in reversed(table.rows):
                cell_value = row.cells[issue_col_idx].text.strip()
                if cell_value and cell_value != "Issue Number" and cell_value != "Issue No":
                    return cell_value
    
    print("Issue number not found, using default: 00")
    return "00"


def format_other_emission_sources(emission_sources, fired_boiler_method=""):
    rows = []
    for src in emission_sources:
        type_name = src.get("type", "").lower()

        if "waste incinerator" in type_name and not INCLUDE_WASTE_INCINERATOR:
            continue

        if any(x in type_name for x in ["main engine", "auxiliary engine", "hydraulic power pack", "boiler"]):
            continue

        row = {
            "ES": src.get("type", "N/A"),
            "METHOD": fired_boiler_method or src.get("method", "N/A")
        }
        rows.append(row)

    return rows


# ---------------- MAIN SCRIPT ----------------
if SELECT_ONE_IMO:
    imos = get_imos_for_company(CSV_FILE, vessel_imo=VESSEL_IMO)
else:
    imos = get_imos_for_company(CSV_FILE, company_name=COMPANY_NAME)
for imo, csv_row in imos.items():
    vessel = get_vessel(imo)
    print(f"Processing vessel with IMO {imo}")
    csv_dwg = csv_row.get("DWG NO.", "UNKNOWN")
    placeholders = format_vessel_placeholder(vessel, csv_dwg)

    if SEEMP_VERSION_1_2:
        process_docx(WORD_TEMPLATE, "temp.docx", placeholders)

        doc = Document("temp.docx")

        emission_sources = get_emission_sources_for_imo(imo)

        fired_boiler_method = get_method_from_placeholder(doc)

        other_es_rows = format_other_emission_sources(emission_sources, fired_boiler_method)

        populate_table(doc, other_es_rows, ["{{ES}}", "{{METHOD}}"])
        
        issue_num = get_issue_number(doc)

        include_bio = has_bio(doc)

        fuel_rows = format_fuel_types(emission_sources, include_bio)

        fuel_placeholders = ["{{TYPE}}", "{{HFO}}", "{{LFO}}", "{{MGO}}"]
        if include_bio:
            fuel_placeholders.append("{{BIO}}")

        populate_table(doc, fuel_rows, fuel_placeholders)

        emis_rows = format_emission_sources(emission_sources, csv_row.get("VERIFIER", ""))
        emis_placeholders = ["{{MODEL}}", "{{DETAILS}}"]
        populate_table(doc, emis_rows, emis_placeholders)

        output_filename = f"{csv_dwg} {vessel['vesselName']} – SEEMP I-II Issue No. {issue_num}"
        doc.save(os.path.join(RESULTS_DIR, f"{output_filename}.docx"))

        if os.path.exists("temp.docx"):
            os.remove("temp.docx")
        
    else:
        doc = Document(WORD_TEMPLATE)
        issue_num = get_issue_number(doc)

        output_filename = f"{csv_dwg} {vessel['vesselName']} – SEEMP PART III Issue No. {issue_num}"
        process_docx(WORD_TEMPLATE, os.path.join(RESULTS_DIR, f"{output_filename}.docx"), placeholders)

    output_doc = os.path.join(RESULTS_DIR, f"{output_filename}.docx")
    output_pdf = os.path.join(RESULTS_DIR, f"{output_filename}.pdf")
    print(f"✅ Saved {output_filename}.docx")
    
    # Convert to PDF
    # convert(output_doc, output_pdf)
    # print(f"✅ Saved {output_filename}.pdf")